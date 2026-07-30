# %%
"""
Build comparison Word document: Pre-revision vs July 2026
Clean academic style — no colored headers.
"""
from pathlib import Path
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import tempfile

BASE    = Path("/Users/asahebi/Library/CloudStorage/GoogleDrive-asahebi@ncsu.edu/My Drive/North Carolina State University/Project - Drug Shortage")
OLD_FIG = BASE / "Paper/Metformin"
NEW_FIG = BASE / "Data/99 - Outputs - Metformin Analysis/processed/outputs"
OUT_DOC = NEW_FIG / "comparison_prerevision_vs_july2026.docx"

# ── helpers ───────────────────────────────────────────────────────────────────
def pdf_to_png(pdf_path, tmp_dir, dpi=200):
    imgs = convert_from_path(str(pdf_path), dpi=dpi, first_page=1, last_page=1)
    out = Path(tmp_dir) / (pdf_path.stem + ".png")
    imgs[0].save(str(out), "PNG")
    return out

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_note(doc, text, size=9):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x50, 0x50, 0x50)
    return p

def write_cell_content(cell, segments, font_size=9):
    """Write formatted runs into a cell. segments: list of {text, bold, red}."""
    cell.text = ''
    p = cell.paragraphs[0]
    for seg in segments:
        text = seg.get('text', '')
        bold = seg.get('bold', False)
        red  = seg.get('red',  False)
        parts = text.split('\n')
        for i, part in enumerate(parts):
            if i > 0:
                run = p.add_run()
                run.add_break()
            if part:
                run = p.add_run(part)
                run.font.size = Pt(font_size)
                if bold:
                    run.bold = True
                if red:
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

def simple_table(doc, headers, rows, col_widths=None, zebra=True):
    """Clean academic table: thin borders, E8E8E8 header, no colored data cells.
    Cell values may be str or list of {text, bold, red} dicts for rich formatting."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'

    hdr = t.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        shade_cell(c, 'E8E8E8')
        c.text = h
        run = c.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for ri, row_data in enumerate(rows):
        tr = t.rows[ri + 1]
        fill = 'F5F5F5' if (zebra and ri % 2 == 1) else 'FFFFFF'
        for ci, val in enumerate(row_data):
            c = tr.cells[ci]
            shade_cell(c, fill)
            if isinstance(val, list):
                write_cell_content(c, val)
            else:
                c.text = str(val)
                if c.paragraphs[0].runs:
                    c.paragraphs[0].runs[0].font.size = Pt(9)

    if col_widths:
        for row in t.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)
    return t

def _add_labeled_figure(doc, label, img, width=6.2):
    p_lbl = doc.add_paragraph()
    r = p_lbl.add_run(label); r.bold = True; r.font.size = Pt(9)
    p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if img and Path(img).exists():
        p_img.add_run().add_picture(str(img), width=Inches(width))
    else:
        p_img.add_run('(not available)')
    doc.add_paragraph()

def add_fig_pair(doc, label_old, old_img, label_new, new_img, width=6.2):
    for label, img in [(label_old, old_img), (label_new, new_img)]:
        _add_labeled_figure(doc, label, img, width)

def add_fig_triple(doc, label_old, old_img, label_new, new_img, label_sfei, sfei_img, width=6.2):
    for label, img in [(label_old, old_img), (label_new, new_img), (label_sfei, sfei_img)]:
        _add_labeled_figure(doc, label, img, width)

SIG_NOTE = '(*) p < 0.05;  (**) p < 0.01;  (***) p < 0.001.  Significant p-values shown in bold red.'

# ── convert old PDFs ──────────────────────────────────────────────────────────
print("Converting old figures...")
with tempfile.TemporaryDirectory() as tmpdir:
    old = {
        1: pdf_to_png(OLD_FIG / "HealthAffairsScholars_Fig1_Price_Volume_by_Inspection.pdf", tmpdir),
        2: pdf_to_png(OLD_FIG / "HealthAffairsScholars_Fig2_Quality_vs_Volume.pdf", tmpdir),
        3: pdf_to_png(OLD_FIG / "HealthAffairsScholars_Fig3_Quality_vs_Price.pdf", tmpdir),
        4: pdf_to_png(OLD_FIG / "HealthAffairsScholars_Fig4_Quality_by_Country.pdf", tmpdir),
    }
    new = {
        1: NEW_FIG / "Figure1_Market_by_Outcome.png",
        2: NEW_FIG / "Figure2_Volume_vs_Quality.png",
        3: NEW_FIG / "Figure3_Price_vs_Quality.png",
        4: NEW_FIG / "Figure4_Quality_by_Country.png",
    }
    new_singlefei = {
        1: NEW_FIG / "Figure1_Market_by_Outcome_SingleFEI.png",
        2: NEW_FIG / "Figure2_Volume_vs_Quality_SingleFEI.png",
        3: NEW_FIG / "Figure3_Price_vs_Quality_SingleFEI.png",
    }
    new_gap36 = {
        1: NEW_FIG / "Figure1_Market_by_Outcome_Gap36.png",
        2: NEW_FIG / "Figure2_Volume_vs_Quality_Gap36.png",
        3: NEW_FIG / "Figure3_Price_vs_Quality_Gap36.png",
    }
    new_sfei_gap36 = {
        1: NEW_FIG / "Figure1_Market_by_Outcome_SingleFEI_Gap36.png",
        2: NEW_FIG / "Figure2_Volume_vs_Quality_SingleFEI_Gap36.png",
        3: NEW_FIG / "Figure3_Price_vs_Quality_SingleFEI_Gap36.png",
    }
    figS1 = NEW_FIG / "FigureS1_Months_Since_Inspection.png"

    # ── build document ─────────────────────────────────────────────────────────
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # ── Title ──────────────────────────────────────────────────────────────────
    title = doc.add_heading('Metformin Analysis: Pre-Revision vs July 2026 Comparison', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_note(doc, 'Prior inspection rule: EventYear strictly < TestYear (same-year inspections excluded).')
    add_note(doc, 'Pre-revision: Health Affairs Scholars, submitted 2026-05-29  |  New pipeline: Steps 1–6, Redica July 2026 refresh')
    doc.add_paragraph()

    # ── Figure 1 ──────────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading('Figure 1 — Market Outcomes by Prior FDA Inspection Outcome', 1)
    add_note(doc, 'Left panel: NADAC price per unit (blank in July 2026 — not yet in pipeline). Right panel: Annual IQVIA extended units by inspection outcome (IND/CHN/USA, log scale).')
    add_note(doc, 'All NDCs = full July 2026 panel including 25 multi-FEI NDC11s.  Single-FEI = those 25 excluded (85 NDC11s, 110 rows).  Both panels restricted to Valisure-tested rows only (145 rows total).')
    add_fig_triple(doc,
        'Pre-revision (Health Affairs Scholars)', old[1],
        'July 2026 — All NDCs (strict rule: EventYear < TestYear)', new[1],
        'July 2026 — Single-FEI NDCs only', new_singlefei[1])

    doc.add_heading('Statistics — Model B (MixedLM + CGM two-way clustered SE)', 2)
    add_note(doc, SIG_NOTE)
    simple_table(doc,
        headers=['Coefficient', 'Pre-revision', 'July 2026 — All NDCs', 'July 2026 — Single-FEI'],
        rows=[
            ['VAI vs NAI',
             [{'text': 'β = −1.820, SE = 0.801\n95% CI [−3.389, −0.250],  '},
              {'text': 'p = 0.025 (*)', 'bold': True, 'red': True}],
             'β = +0.340, SE = 1.538\n95% CI [−2.676, +3.355],  p = 0.826',
             'β = +0.240, SE = 1.554\n95% CI [−2.806, +3.286],  p = 0.878'],
            ['OAI vs NAI',
             'β = +1.747, SE = 0.952\n95% CI [−0.120, +3.613],  p = 0.069',
             'β = +0.902, SE = 2.242\n95% CI [−3.491, +5.296],  p = 0.688',
             'β = −1.248, SE = 1.988\n95% CI [−5.145, +2.649],  p = 0.532'],
            ['OAI vs VAI',
             [{'text': 'β = +3.566, SE = 0.782\n95% CI [+2.033, +5.100],  '},
              {'text': 'p < 0.001 (***)', 'bold': True, 'red': True}],
             'implied ≈ +0.562,  ns',
             'implied ≈ −1.488,  ns'],
        ],
        col_widths=[1.3, 2.3, 2.2, 2.2]
    )
    add_note(doc, (
        'Reference = NAI.  Prior inspection restricted to Drug Quality Assurance outcomes only.  '
        'Panel restricted to Valisure-tested rows only.  '
        'All NDCs: n_obs=106, n_NDC=80, n_FEI=23.  '
        'Single-FEI: n_obs=71, n_NDC=55, n_FEI=18.'
    ))
    doc.add_paragraph()

    doc.add_heading('Descriptive Volume (IQVIA extended units)', 2)
    simple_table(doc,
        headers=['Outcome', 'Pre-rev n', 'Pre-rev Median',
                 'All-NDC n', 'All-NDC Median',
                 'Single-FEI n', 'Single-FEI Median'],
        rows=[
            ['NAI', '39', '15,944,388', '18', '6,174,142',  '18', '6,174,142'],
            ['VAI', '56', '2,392,105',  '77', '2,465,777',  '51', '1,962,991'],
            ['OAI', '16', '18,425,376', '11', '4,357,583',  '2',  '2,697,354'],
        ],
        col_widths=[0.75, 0.65, 1.25, 0.65, 1.25, 0.75, 1.25]
    )
    add_note(doc, 'Pre-revision: Granules India (FEI 3004097901) dominates NAI with 30 NDC-year observations, driving the high NAI mean/median. July 2026: OAI n drops sharply under DQA-only rule (11 all-NDC, 2 single-FEI) because many non-DQA OAI inspections are excluded. Panel restricted to Valisure-tested rows only.')
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('Conclusion (no gap filter): '); r.bold = True; r.font.size = Pt(10)
    p.add_run(
        'No significant relationship between inspection outcome and volume in either the all-NDC '
        '(VAI p = 0.826, OAI p = 0.688, n = 106) or single-FEI (VAI p = 0.878, OAI p = 0.532, n = 71) panel. '
        'OAI cell sizes are very small (11 all-NDC, 2 single-FEI), making OAI estimates highly imprecise.'
    ).font.size = Pt(10)

    # ── Figure 1 — Gap ≤ 36 Months Subsection ──────────────────────────────────
    doc.add_heading('Figure 1 — Gap ≤ 36 Months Subsection (prior inspection ≤ 3 years before test)', 2)
    add_note(doc, 'Excludes rows where the prior DQA inspection is more than 36 months before Jan 1 of the Valisure test year. '
                  'NAI drops from 18 to 8 (long-ago inspections removed); OAI unchanged (11) because all OAI inspections happen to be within 36 months.')
    add_fig_pair(doc,
        'July 2026 — All NDCs (≤36mo gap)', new_gap36[1],
        'July 2026 — Single-FEI (≤36mo gap)', new_sfei_gap36[1])
    simple_table(doc,
        headers=['Coefficient', 'Pre-revision', 'July 2026 — All NDCs ≤36mo', 'July 2026 — Single-FEI ≤36mo'],
        rows=[
            ['VAI vs NAI',
             [{'text': 'β = −1.820, SE = 0.801\n95% CI [−3.389, −0.250],  '},
              {'text': 'p = 0.025 (*)', 'bold': True, 'red': True}],
             'β = +0.152, SE = 2.383\n95% CI [−4.518, +4.822],  p = 0.949',
             'β = +0.046, SE = 2.393\n95% CI [−4.645, +4.737],  p = 0.985'],
            ['OAI vs NAI',
             'β = +1.747, SE = 0.952\n95% CI [−0.120, +3.613],  p = 0.069',
             'β = +0.640, SE = 2.867\n95% CI [−4.979, +6.259],  p = 0.824',
             'β = −1.711, SE = 2.806\n95% CI [−7.211, +3.788],  p = 0.545'],
        ],
        col_widths=[1.3, 2.3, 2.2, 2.2]
    )
    add_note(doc, 'All NDCs ≤36mo: n_obs=76, n_NDC=62, n_FEI=19.  Single-FEI ≤36mo: n_obs=49, n_NDC=41, n_FEI=14.  OAI n=11 (all NDCs) / n=2 (single-FEI).')
    simple_table(doc,
        headers=['Outcome', 'All-NDC ≤36mo n', 'All-NDC ≤36mo Median',
                 'Single-FEI ≤36mo n', 'Single-FEI ≤36mo Median'],
        rows=[
            ['NAI', '8',  '4,659,759',  '8',  '4,659,759'],
            ['VAI', '57', '2,465,777',  '39', '1,606,850'],
            ['OAI', '11', '4,357,583',  '2',  '2,697,354'],
        ],
        col_widths=[0.8, 1.0, 1.5, 1.2, 1.5]
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('Conclusion (≤36mo gap filter): '); r.bold = True; r.font.size = Pt(10)
    p.add_run(
        'After restricting to inspections within 3 years of the test date, NAI drops to n=8 and OAI remains n=11 (all-NDC). '
        'Results remain fully non-significant (VAI p = 0.949, OAI p = 0.824). '
        'The gap filter does not reveal any hidden pattern — if anything, SEs widen due to the smaller NAI cell.'
    ).font.size = Pt(10)

    # ── Figure 2 ──────────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading('Figure 2 — Market Volume vs Tested Drug Quality', 1)
    add_note(doc, 'Each panel pools all available years for that metric. Spearman ρ with NDC-cluster block bootstrap (2,000 resamples).')
    add_note(doc, 'All NDCs = full July 2026 panel (Valisure-tested rows only, 145 rows).  Single-FEI = 25 multi-FEI NDC11s excluded (110 rows).')
    add_fig_triple(doc,
        'Pre-revision (Health Affairs Scholars)', old[2],
        'July 2026 — All NDCs', new[2],
        'July 2026 — Single-FEI NDCs only', new_singlefei[2])

    doc.add_heading('Statistics — Spearman ρ (NDC-cluster block bootstrap, 2,000 resamples)', 2)
    add_note(doc, SIG_NOTE)
    simple_table(doc,
        headers=['Association', 'Pre-revision', 'July 2026 — All NDCs', 'July 2026 — Single-FEI'],
        rows=[
            ['DMF vs Volume',
             [{'text': 'ρ = +0.279,  '},
              {'text': 'p = 0.004 (**)', 'bold': True, 'red': True},
              {'text': '\n95% CI [+0.064, +0.454]'}],
             [{'text': 'ρ = +0.302,  '},
              {'text': 'p_boot = 0.002 (**)', 'bold': True, 'red': True},
              {'text': '\n95% CI [+0.112, +0.467],  n = 126'}],
             [{'text': 'ρ = +0.307,  '},
              {'text': 'p_boot = 0.003 (**)', 'bold': True, 'red': True},
              {'text': '\n95% CI [+0.090, +0.500],  n = 91'}]],
            ['NDMA vs Volume',
             'not significant  (p > 0.10)',
             'ρ = −0.064,  p_boot = 0.635\n95% CI [−0.312, +0.181],  n = 71',
             'ρ = +0.018,  p_boot = 0.898\n95% CI [−0.272, +0.319],  n = 54'],
            ['Diff Factor vs Volume',
             'not significant  (p > 0.10)',
             'ρ = −0.162,  p_boot = 0.454\n95% CI [−0.566, +0.246],  n = 25',
             'ρ = −0.118,  p_boot = 0.613\n95% CI [−0.555, +0.343],  n = 21'],
        ],
        col_widths=[1.5, 1.9, 2.1, 2.1]
    )
    add_note(doc, 'NDC-cluster block bootstrap resamples whole NDC clusters to obtain cluster-robust p-values and 95% CI.')
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('Conclusion (no gap filter): '); r.bold = True; r.font.size = Pt(10)
    p.add_run(
        'Old Observation 2 (DMF vs market volume) is supported and robust in both the all-NDC '
        '(ρ = +0.30, p = 0.002, n = 126) and single-FEI (ρ = +0.31, p = 0.003, n = 91) panels. '
        'NDMA and Difference Factor versus volume remain non-significant.'
    ).font.size = Pt(10)

    # ── Figure 2 — Gap ≤ 36 Months Subsection ──────────────────────────────────
    doc.add_heading('Figure 2 — Gap ≤ 36 Months Subsection (prior inspection ≤ 3 years before test)', 2)
    add_note(doc, 'Same filter as Figure 1 gap section: drops NDC-year rows where prior DQA inspection is >36 months before test year.')
    add_fig_pair(doc,
        'July 2026 — All NDCs (≤36mo gap)', new_gap36[2],
        'July 2026 — Single-FEI (≤36mo gap)', new_sfei_gap36[2])
    simple_table(doc,
        headers=['Association', 'July 2026 — All NDCs (no gap)', 'July 2026 — All NDCs ≤36mo', 'July 2026 — Single-FEI ≤36mo'],
        rows=[
            ['DMF vs Volume',
             [{'text': 'ρ = +0.302,  '},
              {'text': 'p_boot = 0.002 (**)', 'bold': True, 'red': True},
              {'text': '\n95% CI [+0.112, +0.467],  n = 126'}],
             [{'text': 'ρ = +0.359,  '},
              {'text': 'p_boot = 0.001 (**)', 'bold': True, 'red': True},
              {'text': '\n95% CI [+0.152, +0.527],  n = 96'}],
             [{'text': 'ρ = +0.357,  '},
              {'text': 'p_boot = 0.002 (**)', 'bold': True, 'red': True},
              {'text': '\n95% CI [+0.117, +0.551],  n = 69'}]],
            ['NDMA vs Volume',
             'ρ = −0.064,  p_boot = 0.635\n95% CI [−0.312, +0.181],  n = 71',
             'ρ = −0.031,  p_boot = 0.812\n95% CI [−0.263, +0.220],  n = 62',
             'ρ = +0.073,  p_boot = 0.605\n95% CI [−0.212, +0.351],  n = 45'],
            ['Diff Factor vs Volume',
             'ρ = −0.162,  p_boot = 0.454\n95% CI [−0.566, +0.246],  n = 25',
             'ρ = +0.187,  p_boot = 0.546\n95% CI [−0.457, +0.720],  n = 14',
             'ρ = +0.154,  p_boot = 0.638\n95% CI [−0.529, +0.723],  n = 12'],
        ],
        col_widths=[1.4, 2.0, 2.0, 2.0]
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('Conclusion (≤36mo gap filter): '); r.bold = True; r.font.size = Pt(10)
    p.add_run(
        'The DMF–volume relationship strengthens slightly under the gap filter (ρ = +0.36 vs +0.30), '
        'remaining highly significant (p = 0.001). NDMA and Difference Factor remain non-significant. '
        'The gap filter does not weaken the key finding.'
    ).font.size = Pt(10)

    # ── Figure 3 ──────────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading('Figure 3 — Price vs Tested Drug Quality', 1)
    add_note(doc, 'July 2026: Medicaid price per unit = Medicaid Amount Reimbursed / Units Reimbursed (from SDUD). Outliers > $50/unit excluded. Pre-revision used NADAC; 3 NDCs have no SDUD coverage.')
    add_note(doc, 'All NDCs = full July 2026 panel (Valisure-tested rows only).  Single-FEI = 25 multi-FEI NDC11s excluded.')
    add_fig_triple(doc,
        'Pre-revision (Health Affairs Scholars)', old[3],
        'July 2026 — All NDCs (Medicaid price)', new[3],
        'July 2026 — Single-FEI NDCs only (Medicaid price)', new_singlefei[3])

    doc.add_heading('Statistics Comparison — Spearman ρ (NDC-cluster block bootstrap)', 2)
    add_note(doc, SIG_NOTE)
    simple_table(doc,
        headers=['Association', 'Pre-revision (NADAC)', 'July 2026 — All NDCs', 'July 2026 — Single-FEI'],
        rows=[
            ['DMF vs Price',
             'not significant  (p > 0.10)',
             'ρ = −0.110,  p_boot = 0.318\n95% CI [−0.313, +0.128],  n = 117',
             'ρ = −0.115,  p_boot = 0.407\n95% CI [−0.375, +0.155],  n = 84'],
            ['NDMA vs Price',
             [
                 {'text': 'ρ = +0.282,  '},
                 {'text': 'p = 0.013 (*)', 'bold': True, 'red': True},
                 {'text': '\n95% CI [+0.056, +0.490]'},
             ],
             'ρ = −0.091,  p_boot = 0.500\n95% CI [−0.338, +0.187],  n = 68',
             'ρ = −0.161,  p_boot = 0.318\n95% CI [−0.449, +0.160],  n = 51'],
            ['Diff Factor vs Price',
             'not significant  (p > 0.10)',
             'ρ = +0.046,  p_boot = 0.849\n95% CI [−0.399, +0.495],  n = 21',
             'ρ = +0.029,  p_boot = 0.911\n95% CI [−0.470, +0.533],  n = 17'],
        ],
        col_widths=[1.5, 1.9, 1.7, 2.0]
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('Conclusion (no gap filter): '); r.bold = True; r.font.size = Pt(10)
    p.add_run(
        'The pre-revision finding (NDMA vs NADAC price, ρ = +0.282, p = 0.013) does not replicate '
        'with Medicaid price — all associations are non-significant (DMF p = 0.318, NDMA p = 0.500). '
        'This likely reflects the price-source difference (NADAC vs Medicaid reimbursement) rather than '
        'a change in the underlying relationship. NADAC integration remains a future step.'
    ).font.size = Pt(10)

    # ── Figure 3 — Gap ≤ 36 Months Subsection ──────────────────────────────────
    doc.add_heading('Figure 3 — Gap ≤ 36 Months Subsection (prior inspection ≤ 3 years before test)', 2)
    add_note(doc, 'Note: Figure 3 (price vs quality) does not require prior inspection data — the gap filter here drops rows where the NDC-year has a stale inspection (>36mo), which affects only inspection-linked rows in other analyses. All price-quality rows are retained unless the NDC-year is removed from df_gap36.')
    add_fig_pair(doc,
        'July 2026 — All NDCs (≤36mo gap)', new_gap36[3],
        'July 2026 — Single-FEI (≤36mo gap)', new_sfei_gap36[3])
    simple_table(doc,
        headers=['Association', 'July 2026 — All NDCs (no gap)', 'July 2026 — All NDCs ≤36mo', 'July 2026 — Single-FEI ≤36mo'],
        rows=[
            ['DMF vs Price',
             'ρ = −0.110,  p_boot = 0.318\nn = 117',
             'ρ = −0.148,  p_boot = 0.193\nn = 89',
             'ρ = −0.183,  p_boot = 0.197\nn = 63'],
            ['NDMA vs Price',
             'ρ = −0.091,  p_boot = 0.500\nn = 68',
             'ρ = −0.072,  p_boot = 0.622\nn = 59',
             'ρ = −0.150,  p_boot = 0.378\nn = 42'],
            ['Diff Factor vs Price',
             'ρ = +0.046,  p_boot = 0.849\nn = 21',
             'ρ = +0.046,  p_boot = 0.906\nn = 11',
             'ρ = +0.000,  p_boot = 1.000\nn = 9'],
        ],
        col_widths=[1.5, 1.9, 1.7, 2.0]
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('Conclusion (≤36mo gap filter): '); r.bold = True; r.font.size = Pt(10)
    p.add_run(
        'All price vs quality associations remain non-significant after applying the gap filter. '
        'No pattern emerges between Medicaid price and contamination level regardless of inspection recency.'
    ).font.size = Pt(10)

    # ── Figure 4 ──────────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading('Figure 4 — Drug Quality by Country of Manufacture', 1)
    add_note(doc, 'Primary model: MixedLM random NDC intercept + CGM two-way clustered SE (NDC × FEI), reference = USA.')
    add_fig_pair(doc, 'Pre-revision (Health Affairs Scholars)', old[4],
                      'July 2026 (new pipeline)', new[4])

    doc.add_heading('Statistics Comparison — Model B (MixedLM + CGM two-way clustered SE)', 2)
    add_note(doc, SIG_NOTE)
    simple_table(doc,
        headers=['Metric / Comparison', 'Pre-revision', 'July 2026'],
        rows=[
            ['DMF:  IND vs USA',
             'not significant',
             'β = +2.650,  p = 0.136'],
            ['DMF:  CHN vs USA',
             '—',
             'β = +0.277,  p = 0.867'],
            ['DMF:  CHN vs IND',
             '—',
             'β = −2.374,  p = 0.092  (marginal)'],
            ['NDMA:  IND vs USA',
             [
                 {'text': 'β = +1.345,  '},
                 {'text': 'p < 0.001 (***)', 'bold': True, 'red': True},
             ],
             [
                 {'text': 'β = +1.603,  '},
                 {'text': 'p = 0.014 (*)', 'bold': True, 'red': True},
             ]],
            ['NDMA:  CHN vs USA',
             'not significant',
             'β = +0.310,  p = 0.284'],
            ['NDMA:  CHN vs IND',
             [
                 {'text': 'β = −1.090,  '},
                 {'text': 'p = 0.022 (*)', 'bold': True, 'red': True},
             ],
             'β = −1.293,  p = 0.067  (marginal)'],
            ['Diff Factor:  IND vs USA',
             [
                 {'text': 'β = +0.117,  '},
                 {'text': 'p = 0.011 (*)', 'bold': True, 'red': True},
             ],
             'β = +0.074,  p = 0.061  (marginal)'],
            ['Diff Factor:  CHN vs USA',
             'not significant',
             'β = +0.060,  p = 0.185'],
            ['Diff Factor:  CHN vs IND',
             'not significant',
             'β = −0.015,  p = 0.802'],
        ],
        col_widths=[2.0, 2.5, 2.6]
    )
    add_note(doc, (
        'Descriptive means (July 2026): '
        'DMF — IND 28,607 ng/day, CHN 3,355, USA 4,696.  '
        'NDMA — IND 65.2 ng/day, CHN 2.0, USA 0.0.  '
        'Difference Factor — IND 0.261, CHN 0.226, USA 0.153.'
    ))
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('Conclusion: '); r.bold = True; r.font.size = Pt(10)
    p.add_run(
        'Old Observation 3 (NDMA, India vs USA) remains statistically supported (p = 0.014, vs p < 0.001 previously). '
        'The claims that NDMA is lower in China than India (old p = 0.022, new p = 0.067) and that '
        'dissolution failure is more common in India than the US (old p = 0.011, new p = 0.061) '
        'are weakened to marginal significance and should be softened in the paper. '
        'DMF country differences remain non-significant in both datasets.'
    ).font.size = Pt(10)

    # ── Figure S1 — Months Since Last Inspection ──────────────────────────────
    doc.add_page_break()
    doc.add_heading('Figure S1 — Distribution of Months Since Last Inspection', 1)
    add_note(doc, (
        'For each (NDC11 × TestYear) row with a prior classified inspection, '
        '"months since inspection" = exact days from inspection end date to Jan 1 of test year ÷ (365.25/12). '
        'Uses exact inspection end dates from Redica (Event End Date column). '
        'This shows how stale the prior inspection signal is relative to the Valisure test year.'
    ))
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if figS1.exists():
        p_img.add_run().add_picture(str(figS1), width=Inches(6.2))
    else:
        p_img.add_run('(FigureS1_Months_Since_Inspection.png not found)')
    doc.add_paragraph()

    doc.save(str(OUT_DOC))
    print(f'Saved: {OUT_DOC}')

print('Done.')
# %%
